import os
import re
import uuid
import logging
from datetime import datetime
from functools import wraps
from io import BytesIO

from dotenv import load_dotenv
from supabase import create_client, Client
from flask import (
    Flask, render_template, request, redirect, url_for, 
    session, flash, send_file
)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Flowable
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import RequestEntityTooLarge

from utils.speech_to_text import convert_to_text
from utils.ai_summarizer import generate_output

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = (os.getenv("FLASK_SECRET_KEY") or "default-dev-key").strip()

UPLOAD_FOLDER = os.path.join("static", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB Limit

AUDIO_EXTENSIONS = {"mp3", "wav", "m4a", "flac", "ogg"}
VIDEO_EXTENSIONS = {"mp4", "avi", "mov", "mkv", "webm"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}
DOCUMENT_EXTENSIONS = {"pdf", "docx", "txt"}

SUPABASE_URL = os.getenv("SUPABASE_URL", "").strip()
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "").strip()
if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError("SUPABASE_URL and SUPABASE_KEY must be set in the .env file")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

def allowed_file(filename, allowed_set):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_set

def get_clean_filename(output_text, extension):
    if not output_text:
        return f"document.{extension}"
    first_line = output_text.split('\n')[0].replace('*', '')
    clean_name = re.sub(r'[\\/*?:"<>|]', "", first_line).strip()
    return f"{clean_name[:50] or 'output'}.{extension}"

def extract_text_from_document(file_path, ext):
    if ext == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == "docx":
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == "pdf":
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    return ""

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated_function

@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    flash("File size exceeds the 500MB upload limit.")
    return redirect(url_for('dashboard'))

@app.route("/")
def landing():
    return render_template("land.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
        
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        first = (request.form.get("firstname") or "").strip()
        last = (request.form.get("lastname") or "").strip()
        email = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "").strip()
        date = (request.form.get("date") or "").strip()

        if not all([username, first, last, email, password, date]):
            flash("All registration fields are required.")
            return redirect(url_for("register"))

        if len(password) < 8:
            flash("Password must be at least 8 characters long.")
            return redirect(url_for("register"))
        
        try:
            existing = supabase.table("users").select("id").or_(f"username.eq.{username},email.eq.{email}").execute()
            if existing.data:
                flash("Username or Email already exists.")
                return redirect(url_for("register"))

            user_data = {
                "username": username,
                "first_name": first,
                "last_name": last,
                "email": email,
                "password_hash": generate_password_hash(password),
                "date_of_birth": date
            }
            supabase.table("users").insert(user_data).execute()
            flash("Registration successful! Please sign in.")
            return redirect(url_for("login"))
        except Exception as e:
            logger.error(f"Registration Error: {e}")
            flash("An error occurred during registration.")
            return redirect(url_for("register"))

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
        
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        password = (request.form.get("password") or "").strip()

        if not username or not password:
            flash("All fields are required.")
            return redirect(url_for("login"))
             
        try:
            response = supabase.table("users").select("*").eq("username", username).execute()
            if response.data and len(response.data) > 0:
                user = response.data[0]
                if check_password_hash(str(user.get("password_hash") or ""), password):
                    session.clear()
                    session.update({
                        "user_id": user.get("id"),
                        "username": user.get("username"),
                        "first_name": user.get("first_name"),
                        "last_name": user.get("last_name"),
                        "email": user.get("email"),
                        "date": user.get("date_of_birth"),
                        "profile_image": user.get("profile_image")
                    })
                    return redirect(url_for("dashboard"))
            flash("Invalid username or password.")
        except Exception as e:
            logger.error(f"Login Error: {e}")
            flash("Login failed due to system error.")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("You have been logged out.")
    response = redirect(url_for("landing"))
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    return response

@app.route("/dashboard")
@login_required
def dashboard():
    last_id = session.pop("last_history_id", None)
    selected = session.pop("selected", "notes") 
    result = None
    input_text = None

    if last_id:
        try:
            response = supabase.table("history").select("input_text, output_text").eq("id", last_id).execute()
            if response.data and len(response.data) > 0:
                record = response.data[0]
                result = str(record.get("output_text") or "")
                input_text = str(record.get("input_text") or "")
        except Exception as e:
            logger.error(f"Error fetching dashboard data: {e}")
            
    return render_template("index.html", result=result, input_text=input_text, selected=selected)

@app.route("/upload", methods=["POST"])
@login_required
def upload():  
    output_type = request.form.get("output_type", "notes")
    user_prompt = (request.form.get("user_prompt") or "").strip()
    file = request.files.get("audio_file")
    
    text = None
    input_type = None
    path = None 

    try:
        if file and file.filename:
            filename = secure_filename(file.filename)
            unique_name = f"{uuid.uuid4().hex}_{filename}"
            path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
            file.save(path) 

            ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
            
            # FIXED: All document formats (PDF, DOCX, TXT) properly extracted
            if ext in DOCUMENT_EXTENSIONS:
                text = extract_text_from_document(path, ext)
                input_type = "document"
            elif ext in AUDIO_EXTENSIONS or ext in VIDEO_EXTENSIONS:
                try:
                    text = convert_to_text(path)
                except RuntimeError:
                    flash("AI speech model is loading, please retry in 10 seconds.")
                    return redirect(url_for("dashboard"))
                input_type = "video" if ext in VIDEO_EXTENSIONS else "audio"
            else:
                flash("Unsupported file format uploaded.")
                return redirect(url_for("dashboard"))
                
        elif user_prompt:
            text = user_prompt
            input_type = "prompt"
        else:
            flash("Please paste text or choose a file.")
            return redirect(url_for("dashboard"))
            
        if not text or not text.strip():
            raise ValueError("Could not extract readable text from the provided input.")

        result = generate_output(text, output_type)
        if result and ("error" in str(result).lower()):
            flash(f"AI Generation Issue: {result}")
            return redirect(url_for("dashboard"))

        history_data = {
            "user_id": session["user_id"],
            "input_type": input_type,
            "input_text": text,
            "output_text": result,
            "output_type": output_type
        }
        
        response = supabase.table("history").insert(history_data).execute()
        if response.data and len(response.data) > 0:
            session["last_history_id"] = response.data[0].get('id')
            session["selected"] = output_type
        
        return redirect(url_for("dashboard"))

    except Exception as e:
        logger.error(f"Processing error: {e}")
        flash(f"Error processing content: {str(e)[:100]}")
        return redirect(url_for("dashboard"))
    finally:
        if path and os.path.exists(path):
            os.remove(path)

@app.route("/history")
@login_required
def history():
    query = request.args.get("q", "").strip()
    try:
        db_query = supabase.table("history").select("*").eq("user_id", session["user_id"]).order("created_at", desc=True)
        if query:
            db_query = db_query.or_(f"input_text.ilike.%{query}%,output_text.ilike.%{query}%")
        records = db_query.execute().data or []
        for record in records:
            raw_date = record.get("created_at")
            try:
                dt_obj = datetime.fromisoformat(raw_date.replace('Z', '+00:00'))
                record["formatted_date"] = dt_obj.strftime('%Y-%m-%d')
            except Exception:
                record["formatted_date"] = "N/A"
    except Exception as e:
        logger.error(f"History Fetch Error: {e}")
        records = []
    return render_template("history.html", records=records, query=query)

@app.route("/history/delete/<int:history_id>", methods=["POST"])
@login_required
def delete_history_item(history_id):
    supabase.table("history").delete().eq("id", history_id).eq("user_id", session["user_id"]).execute()
    return redirect(url_for("history"))

@app.route("/history/delete_all", methods=["POST"])
@login_required
def delete_all_history():
    supabase.table("history").delete().eq("user_id", session["user_id"]).execute()
    return redirect(url_for("history"))

@app.route("/history/<int:history_id>/pdf")
@login_required
def download_history_pdf(history_id):
    response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
    if not response.data: return redirect(url_for("history"))
    record = response.data[0]
    output_text = str(record.get("output_text") or "")
    output_type = str(record.get("output_type") or "output").replace("_", " ").title()

    fname = get_clean_filename(output_text, "pdf")
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{output_type}</b><br/><br/>" + output_text.replace("\n", "<br/>"), styles["Normal"])]
    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="application/pdf")

@app.route("/history/<int:history_id>/docx")
@login_required
def download_history_docx(history_id):
    response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
    if not response.data: return redirect(url_for("history"))
    record = response.data[0]
    output_text = str(record.get("output_text") or "")
    output_type = str(record.get("output_type") or "output").replace("_", " ").title()

    fname = get_clean_filename(output_text, "docx")
    doc = Document()
    doc.add_heading(output_type, level=1)
    doc.add_paragraph(output_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/history/<int:history_id>/txt")
@login_required
def download_history_txt(history_id):
    response = supabase.table("history").select("*").eq("id", history_id).eq("user_id", session["user_id"]).execute()
    if not response.data: return redirect(url_for("history"))
    record = response.data[0]
    output_text = str(record.get("output_text") or "")
    output_type = str(record.get("output_type") or "output").replace("_", " ").upper()

    fname = get_clean_filename(output_text, "txt")
    buffer = BytesIO(f"{output_type}:\n\n{output_text}".encode('utf-8'))
    return send_file(buffer, as_attachment=True, download_name=fname, mimetype="text/plain")

@app.route("/profile")
@login_required
def profile():
    dob = session.get("date")
    formatted_date = dob
    if dob and isinstance(dob, str):
        try:
            formatted_date = datetime.strptime(dob, "%Y-%m-%d").strftime("%d %B %Y")
        except ValueError:
            pass 
    return render_template("profile.html", username=session.get("username"), firstname=session.get("first_name"), lastname=session.get("last_name"), email=session.get("email"), date=formatted_date, profile_image=session.get("profile_image"))

@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    if request.method == "POST":
        current_password = (request.form.get("current_password") or "").strip()
        new_password = (request.form.get("new_password") or "").strip()
        confirm_password = (request.form.get("confirm_password") or "").strip()
        
        if len(new_password) < 8:
            flash("New password must be at least 8 characters long.")
            return redirect(url_for("change_password"))
        if new_password != confirm_password:
            flash("New passwords do not match.")
            return redirect(url_for("change_password"))
            
        try:
            response = supabase.table("users").select("password_hash").eq("id", session["user_id"]).execute()
            if not response.data or not check_password_hash(str(response.data[0].get("password_hash") or ""), current_password):
                flash("Current password is incorrect.")
                return redirect(url_for("change_password"))
            
            supabase.table("users").update({"password_hash": generate_password_hash(new_password)}).eq("id", session["user_id"]).execute()
            flash("Password changed successfully.")
            return redirect(url_for("profile"))
        except Exception as e:
            logger.error(f"Password Change Error: {e}")
            flash("Error changing password.")

    return render_template("change_password.html")

@app.route("/upload-profile-photo", methods=["POST"])
@login_required
def upload_profile_photo():
    file = request.files.get("profile_image")
    if not file or not file.filename or not allowed_file(file.filename, IMAGE_EXTENSIONS):
        flash("Invalid image selected.")
        return redirect(url_for("profile"))
        
    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
    file.save(path)
    
    try:
        old_data = supabase.table("users").select("profile_image").eq("id", session["user_id"]).execute()
        old_image = old_data.data[0].get("profile_image") if old_data.data else None
        
        supabase.table("users").update({"profile_image": unique_name}).eq("id", session["user_id"]).execute()
        if old_image:
            old_path = os.path.join(app.config["UPLOAD_FOLDER"], old_image)
            if os.path.exists(old_path): os.remove(old_path)
                
        session["profile_image"] = unique_name
        flash("Profile photo updated successfully.")
    except Exception as e:
        logger.error(f"Photo Upload Error: {e}")
        flash("Error updating profile photo.")
    return redirect(url_for("profile"))

@app.route("/delete-profile-photo", methods=["POST"])
@login_required
def delete_profile_photo():
    try:
        old_data = supabase.table("users").select("profile_image").eq("id", session["user_id"]).execute()
        old_image = old_data.data[0].get("profile_image") if old_data.data else None
        if old_image:
            supabase.table("users").update({"profile_image": None}).eq("id", session["user_id"]).execute()
            path = os.path.join(app.config["UPLOAD_FOLDER"], old_image)
            if os.path.exists(path): os.remove(path)
            session["profile_image"] = None
            flash("Profile photo deleted.")
    except Exception as e:
        logger.error(f"Photo Delete Error: {e}")
    return redirect(url_for("profile"))

@app.route("/delete-account", methods=["POST"])
@login_required
def delete_account():
    user_id = session["user_id"]
    try:
        old_data = supabase.table("users").select("profile_image").eq("id", user_id).execute()
        image = old_data.data[0].get("profile_image") if old_data.data else None
        supabase.table("users").delete().eq("id", user_id).execute()
        if image:
            img_path = os.path.join(app.config["UPLOAD_FOLDER"], image)
            if os.path.exists(img_path): os.remove(img_path)
        session.clear()
        flash("Your account has been permanently deleted.")
    except Exception as e:
        logger.error(f"Account Delete Error: {e}")
        return redirect(url_for("profile"))
    return redirect(url_for("landing"))

@app.after_request
def add_header(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    return response

if __name__ == "__main__":
    app.run(debug=True)